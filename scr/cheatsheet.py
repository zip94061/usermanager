import os.path
import os
import platform
import subprocess
import logging

from pathlib import Path

from docxtpl import DocxTemplate
from docx import Document


class CheatSheet:
    def __init__(self, full_name, login, email, password, phone):
        self.logger = logging.getLogger(__name__)
        self.message = ''
        self.full_name = full_name
        self.login = login
        self.email = email
        self.password = password
        self.phone = phone
        self.template_file = Path("template", "template.docx")
        if not self.template_file.exists():
            self.template_file.parent.mkdir(parents=True, exist_ok=True)
            self.create_default_template()
            self.message = (f'Файл шаблона не существует. Создан шаблон поумолчанию'
                            )
            self.logger.info(self.message)
            self.logger.info(f'\nВ шаблоне должны быть поля:\n'
                             f'{{ {{ full_name }} }}\n'
                             f'{{ {{ login }} }}\n'
                             f'{{ {{ email }} }}\n'
                             f'{{ {{ password }} }}\n'
                             f'{{ {{ phone }} }}\n')
        self.template = DocxTemplate(self.template_file)
        self.user_file = Path("temp", "new_user.docx")
        self.user_file.parent.mkdir(parents=True, exist_ok=True)

    def set_data(self):
        text = {
            "full_name" : self.full_name,
            "login" : self.login,
            "email" : self.email,
            "password" : self.password,
            "phone" : self.phone,
            }
        return text

    def write_data(self):
        self.template.render(self.set_data())
        self.template.save(self.user_file)

    def make_cheat_sheet(self):
        self.set_data()
        self.write_data()
        current_os = platform.system()
        if current_os == 'Linux':
            subprocess.run(['open', self.user_file])
        if current_os == 'Windows':
            os.startfile(self.user_file)

    def get_message(self):
        return self.message

    def create_default_template(self):
        doc = Document()

        doc.add_heading('{{ full_name }}', level=0)
        doc.add_paragraph('"login": {{ login }}')
        doc.add_paragraph('"email": {{ email }}')
        doc.add_paragraph('"password": {{ password }}')
        doc.add_paragraph('"phone": {{ phone }}')

        doc.save(str(self.template_file))
